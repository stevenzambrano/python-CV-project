#installed " pip3 install python-docx"
#pip3 install --upgrade pip // this to update
from docx import Document
import docx 
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('Nico go fuck your self')

##NOW WE WILL WRITE INSIDE THE DOCUMENT!
document = Document()

document.add_picture('IMG_3067.JPG', width= Inches(2.0))  ## this is the code to add a picture, copy the exact extention of the file


#name = 'Steven'
#phone_number = '07881246298'
#email='stevenzambrano@hotmail.co.uk'

name = input('What is your name fatty? ')
phone_number = input('What is your email ? ')
email= input('What is your email cunt ?')


# adding the above 3 variables into our CV/document
document.add_paragraph(
    name + ' | ' + phone_number + ' | '+ email)

#About me
document.add_heading('About me') # becareful with the fullstop and underscore
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

#Work experience
document.add_heading('Work Experience')
work_experience = input('I work for VSL systems')
document.add_paragraph(work_experience)

p = document.add_paragraph()
company = input('Enter company ?')
from_date = input('From date ? ')
to_date = input('To date ?')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at' + company)
p.add_run(experience_details)

#more experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No')
    if has_more_experiences.lower() == 'yes':# this means the answer will be written in lowercase
        p = document.add_paragraph()
                
        company = input('Enter company ?')
        from_date = input('From date ? ')
        to_date = input('To date ?')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at' + company+ ' ')
        p.add_run(experience_details)
    else:
        break
        
#NEW SKILLS

document.add_heading('Skills')
Skills = input('I am a programmer')
document.add_paragraph(Skills)
document.add_paragraph(Skills).style = 'List Bullet'

x = document.add_paragraph()
skill_set = input('what skills do you have')

while True:
    has_more_skills = input('Do you have more skills? Yes or No')
    if has_more_skills.lower() == 'yes':
        first_skill = input('describe')
        from_date = input('From date?')
        to_date = input('to date?')

        x.add_run(first_skill + ' ').bold = True
        x.add_run(from_date + ' '+ to_date +'\n').italic = True
    else:
        break


document.save('cv.docx')

#we wrote all our functions we needed from requirements text and installed it in the bottom of our terminal
#pip3 install -r requirements.txt